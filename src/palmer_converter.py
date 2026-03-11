"""palmer_converter.py – Replace \\Palmer commands in a Word document with rendered images.

Provides the core conversion logic used by the GUI Converter tab.
Independent of tkinter so that it can be tested or used standalone.

Requires: python-docx, Pillow  (imported at module level)
"""

from __future__ import annotations

import functools
import itertools
import logging
import re
import subprocess
import tempfile
import threading
from copy import deepcopy
from pathlib import Path
from typing import Callable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from PIL import Image

from palmer_engine import PalmerCompiler, DEFAULT_FONT_FAMILY, DEFAULT_FONT_SIZE_PT

logger = logging.getLogger(__name__)


class ConversionCancelled(Exception):
    """Raised when a ``stop_event`` signals cancellation inside :func:`convert_docx`."""


# English Metric Units per inch — used to convert pixel dimensions to the
# EMU values that python-docx requires for inline image sizing.
EMU_PER_INCH = 914_400

# ---------------------------------------------------------------------------
# CJK font detection — used by _extract_font() to prefer the eastAsia font
# when the ASCII font is a generic Word default.
# ---------------------------------------------------------------------------

# ASCII font names that Word commonly assigns as defaults in CJK documents.
_GENERIC_ASCII_FONTS: frozenset[str] = frozenset({
    "Times New Roman",
    "Century",
    "Calibri",
    "Arial",
    "Cambria",
})


@functools.cache
def _get_system_fonts() -> frozenset[str]:
    """Return the set of font family names installed on the system.

    Uses ``fc-list`` (fontconfig), which is available on any system where
    XeLaTeX works.  The result is cached for the lifetime of the process.
    Returns an empty set if ``fc-list`` is unavailable.
    """
    try:
        result = subprocess.run(
            ["fc-list", "--format", "%{family}\n"],
            capture_output=True, text=True, timeout=10,
        )
        if result.returncode != 0:
            return frozenset()
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError) as exc:
        logger.debug("fc-list unavailable: %s", exc)
        return frozenset()
    if not result.stdout:
        logger.debug("fc-list returned empty output")
        return frozenset()
    names: set[str] = set()
    for line in result.stdout.splitlines():
        # fc-list may return comma-separated aliases, e.g. "Yu Mincho,游明朝"
        for name in line.split(","):
            stripped = name.strip()
            if stripped:
                names.add(stripped)
    logger.debug("fc-list detected %d font families", len(names))
    return frozenset(names)


def _is_system_font(name: str) -> bool:
    """Return True if *name* is installed on the system (via fontconfig)."""
    if not name:
        return False
    return name in _get_system_fonts()


# ---------------------------------------------------------------------------
# Theme font resolution — Word documents often specify fonts via theme
# references (e.g. minorHAnsi, minorEastAsia) rather than literal names.
# python-docx does not resolve these, so we do it ourselves.
# ---------------------------------------------------------------------------

_DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_THEME_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
)


@functools.cache
def _resolve_theme_fonts(blob: bytes) -> dict[str, str]:
    """Parse a theme XML *blob* and return a mapping of theme-slot to font name.

    The returned dict maps keys like ``"minorHAnsi"``, ``"minorEastAsia"``,
    ``"majorHAnsi"``, ``"majorEastAsia"`` to the concrete font family name
    defined in the theme.

    *blob* is ``bytes`` (immutable, hashable) and is used directly as the
    cache key.  This avoids the instability of ``id(part)``-based keys, where
    garbage-collected objects can have their IDs reused.
    """
    from lxml import etree  # type: ignore[import-untyped]  # transitive dep of python-docx

    mapping: dict[str, str] = {}
    try:
        root = etree.fromstring(blob)
    except etree.XMLSyntaxError:
        return mapping

    ns = {"a": _DRAWINGML_NS}

    for prefix, xpath in (("minor", ".//a:minorFont"), ("major", ".//a:majorFont")):
        group = root.find(xpath, ns)
        if group is None:
            continue
        # Latin (HAnsi) typeface
        latin = group.find("a:latin", ns)
        if latin is not None and latin.get("typeface"):
            mapping[f"{prefix}HAnsi"] = latin.get("typeface")
        # eastAsia typeface — prefer the script-specific <a:font script="Jpan">
        # entry, which overrides the generic <a:ea> for Japanese documents.
        # Fall back to <a:ea> if no Jpan entry exists.
        jpan = group.find('a:font[@script="Jpan"]', ns)
        ea = group.find("a:ea", ns)
        ea_name = ""
        if jpan is not None and jpan.get("typeface"):
            ea_name = jpan.get("typeface")
        elif ea is not None and ea.get("typeface"):
            ea_name = ea.get("typeface")
        if ea_name:
            mapping[f"{prefix}EastAsia"] = ea_name
    logger.debug("Theme fonts resolved: %s", mapping)
    return mapping


def _get_theme_fonts(run) -> dict[str, str]:
    """Return the theme-font mapping for the document that *run* belongs to.

    Returns an empty dict when the theme cannot be found.
    """
    try:
        part = run.part
        for _key, rel in part.rels.items():
            if rel.reltype == _THEME_RELTYPE:
                blob = rel.target_part.blob
                return _resolve_theme_fonts(blob)
    except (AttributeError, KeyError) as exc:
        logger.debug("Theme font resolution failed: %s", exc)
    return {}


# ---------------------------------------------------------------------------
# Palmer command parser
# ---------------------------------------------------------------------------

def _read_braced(text: str, pos: int) -> tuple[str | None, int]:
    """Read content between matched braces starting at *pos*.

    Returns ``(content, next_pos)`` or ``(None, pos)`` on failure.
    Escaped braces (``\\{`` and ``\\}``) are treated as literal characters
    and do not affect the depth count, matching the behaviour of
    ``_check_brace_balance`` in ``palmer_engine.py``.
    """
    if pos >= len(text) or text[pos] != "{":
        return None, pos
    depth = 1
    start = pos + 1
    pos += 1
    while pos < len(text) and depth > 0:
        if text[pos] == "\\" and pos + 1 < len(text) and text[pos + 1] in "{}":
            pos += 2  # skip escaped brace
            continue
        if text[pos] == "{":
            depth += 1
        elif text[pos] == "}":
            depth -= 1
        pos += 1
    if depth != 0:
        return None, pos
    return text[start : pos - 1], pos


def _normalize_yen_sign(text: str) -> str:
    r"""Replace the fullwidth yen sign (U+00A5) with a backslash (U+005C).

    On Japanese Windows the yen sign ``\u00a5`` is commonly displayed in place
    of the backslash, and users may type ``\u00a5Palmer`` instead of
    ``\\Palmer``.  Normalising early keeps the rest of the parser simple.
    """
    return text.replace("\u00a5", "\\")


def find_palmer_commands(text: str) -> list[dict]:
    r"""Return a list of dicts for each ``\Palmer`` command found in *text*.

    Both backslash (``\\``) and the yen sign (``\u00a5``) are accepted as the
    command prefix to support Japanese Windows environments.

    Each dict contains keys: start, end, option, UL, UR, LR, LL,
    upper_mid, lower_mid.  ``start`` / ``end`` refer to positions in the
    **original** (un-normalised) text.
    """
    # Normalise yen signs so that the parser only needs to look for backslash.
    norm = _normalize_yen_sign(text)
    TAG = "\\Palmer"
    results: list[dict] = []
    i = 0
    while i < len(norm):
        idx = norm.find(TAG, i)
        if idx == -1:
            break
        pos = idx + len(TAG)

        # Optional [option]
        option: str | None = None
        if pos < len(norm) and norm[pos] == "[":
            close = norm.find("]", pos)
            if close == -1:
                i = pos
                continue
            option = norm[pos + 1 : close]
            pos = close + 1

        # 6 mandatory {…} arguments
        args: list[str] = []
        ok = True
        for _ in range(6):
            while pos < len(norm) and norm[pos] in " \t\n\r":
                pos += 1
            arg, pos = _read_braced(norm, pos)
            if arg is None:
                ok = False
                break
            args.append(arg)

        if ok and len(args) == 6:
            results.append(
                {
                    "start": idx,
                    "end": pos,
                    "option": option,
                    "UL": args[0],
                    "UR": args[1],
                    "LR": args[2],
                    "LL": args[3],
                    "upper_mid": args[4],
                    "lower_mid": args[5],
                }
            )
            i = pos
        else:
            i = idx + 1
    return results


# ---------------------------------------------------------------------------
# Alt-text generation for accessibility
# ---------------------------------------------------------------------------

# Valid alt-text modes.
ALT_TEXT_MODES = (None, "FDI", "Universal", "Anatomical", "Alphanumeric", "Palmer command")

# Inline vertical-alignment modes for the Converter tab.
VALIGN_MODES = ("Force center", "Follow command option")

_VALID_TOOTH_CHARS = set("12345678ABCDE")

_ANATOMICAL_NAMES: dict[str, str] = {
    "1": "Central incisor",
    "2": "Lateral incisor",
    "3": "Canine",
    "4": "First premolar",
    "5": "Second premolar",
    "6": "First molar",
    "7": "Second molar",
    "8": "Third molar",
    "A": "Primary Central incisor",
    "B": "Primary Lateral incisor",
    "C": "Primary Canine",
    "D": "Primary First molar",
    "E": "Primary Second molar",
}

# FDI quadrant numbers: permanent (1-4) and deciduous (5-8).
_FDI_QUAD: dict[str, tuple[int, int]] = {
    "UL": (1, 5),  # patient's right upper
    "UR": (2, 6),  # patient's left upper
    "LR": (3, 7),  # patient's left lower
    "LL": (4, 8),  # patient's right lower
}

# Alphanumeric quadrant prefixes (patient's perspective).
# Engine keys use viewer/dentist perspective; Alphanumeric notation uses
# patient's perspective, so left and right are swapped.
_ALPHA_QUAD: dict[str, str] = {
    "UL": "UR",   # engine UL (viewer's left) = patient's upper right
    "UR": "UL",   # engine UR (viewer's right) = patient's upper left
    "LR": "LL",   # engine LR (viewer's right) = patient's lower left
    "LL": "LR",   # engine LL (viewer's left) = patient's lower right
}

# Anatomical quadrant labels.
_ANAT_QUAD: dict[str, str] = {
    "UL": "Right Maxillary",
    "UR": "Left Maxillary",
    "LR": "Left Mandibular",
    "LL": "Right Mandibular",
}

# Universal Numbering System mapping.
# Maps (quadrant_key, palmer_tooth_char) → Universal number/letter.
_UNIVERSAL_PERM: dict[str, dict[str, int]] = {
    "UL": {str(i): 9 - i for i in range(1, 9)},      # 1→8, 2→7, ...8→1
    "UR": {str(i): 8 + i for i in range(1, 9)},       # 1→9, 2→10,...8→16
    "LR": {str(i): 25 - i for i in range(1, 9)},      # 1→24,2→23,...8→17
    "LL": {str(i): 24 + i for i in range(1, 9)},      # 1→25,2→26,...8→32
}
_UNIVERSAL_DECID: dict[str, dict[str, str]] = {
    "UL": {"A": "E", "B": "D", "C": "C", "D": "B", "E": "A"},
    "UR": {"A": "F", "B": "G", "C": "H", "D": "I", "E": "J"},
    "LR": {"A": "O", "B": "N", "C": "M", "D": "L", "E": "K"},
    "LL": {"A": "P", "B": "Q", "C": "R", "D": "S", "E": "T"},
}


_RE_TEX_CMD_BRACE = re.compile(r"\\[a-zA-Z]+\{([^}]*)\}")
_RE_BARE_TEX_CMD = re.compile(r"\\[a-zA-Z]+")
_RE_BRACES_SPACES = re.compile(r"[{}\s]")


def _strip_tex_commands(value: str) -> str:
    r"""Remove TeX formatting commands and return only the content characters.

    Examples: ``\textbf{1}2`` → ``12``, ``\underline{AB}`` → ``AB``.
    """
    # Iteratively unwrap \command{content} → content.
    prev = None
    text = value
    while text != prev:
        prev = text
        text = _RE_TEX_CMD_BRACE.sub(r"\1", text)
    # Remove any remaining bare \commands (e.g. \relax).
    text = _RE_BARE_TEX_CMD.sub("", text)
    # Strip spaces, braces.
    text = _RE_BRACES_SPACES.sub("", text)
    return text


_RE_DIGIT_RANGE = re.compile(r"([1-8])-([1-8])")
_RE_LETTER_RANGE = re.compile(r"([A-E])-([A-E])")


def _expand_ranges(text: str) -> str:
    r"""Expand range shorthand in a quadrant field.

    ``1-4`` → ``1234``, ``A-C`` → ``ABC``.  Multiple ranges and bare
    characters can coexist: ``1-3A-C`` → ``123ABC``.
    """

    def _repl(m: re.Match) -> str:  # noqa: D401
        a, b = ord(m.group(1)), ord(m.group(2))
        if a > b:
            raise ValueError(
                f"Invalid range '{m.group(0)}': start ({m.group(1)}) "
                f"must not be greater than end ({m.group(2)})"
            )
        return "".join(chr(c) for c in range(a, b + 1))

    text = _RE_DIGIT_RANGE.sub(_repl, text)
    text = _RE_LETTER_RANGE.sub(_repl, text)
    return text


def _expand_count_from_midline(value: str) -> str:
    """Expand tooth notation to include the mesial tooth when a midline dash
    is present.

    When the midline field contains a dash (``-``), the most mesial tooth
    (``1`` for permanent, ``A`` for deciduous) is implied by the dash.

    **Single character** — interpreted as the outermost tooth; all teeth from
    the midline to that tooth are generated:

        ``"2"`` → ``"12"``, ``"3"`` → ``"123"``.
        ``"B"`` → ``"AB"``, ``"C"`` → ``"ABC"``.

    **Multiple characters** — treated as explicitly listed teeth; consecutive
    teeth from the midline up to (but not including) the most mesial listed
    tooth are prepended.  Gaps between listed teeth are NOT filled:

        ``"BC"`` → ``"ABC"``, ``"BE"`` → ``"ABE"``.
        ``"46"`` → ``"12346"``, ``"358"`` → ``"12358"``.
        ``"245"`` → ``"1245"``, ``"13"`` → ``"13"`` (already contains 1).

    Values with invalid characters are returned unchanged (they will later
    fail validation).
    """
    if len(value) == 1:
        if value in "12345678":
            return "".join(str(i) for i in range(1, int(value) + 1))
        if value in "ABCDE":
            return "".join(
                chr(c) for c in range(ord("A"), ord(value) + 1)
            )
    elif len(value) > 1:
        if all(c in "12345678" for c in value):
            m = min(int(c) for c in value)
            prefix = "".join(str(i) for i in range(1, m))
            return prefix + value if prefix else value
        if all(c in "ABCDE" for c in value):
            m = min(ord(c) for c in value)
            prefix = "".join(chr(c) for c in range(ord("A"), m))
            return prefix + value if prefix else value
    return value


_MIDLINE_DASHES = frozenset((
    "\u002D",  # HYPHEN-MINUS
    "\u2010",  # HYPHEN
    "\u2011",  # NON-BREAKING HYPHEN
    "\u2012",  # FIGURE DASH
    "\u2013",  # EN DASH
    "\u2014",  # EM DASH
    "\u2015",  # HORIZONTAL BAR
    "\u2212",  # MINUS SIGN
    "\u207B",  # SUPERSCRIPT MINUS
    "\u208B",  # SUBSCRIPT MINUS
    "\uFE58",  # SMALL EM DASH
    "\uFE63",  # SMALL HYPHEN-MINUS
    "\uFF0D",  # FULLWIDTH HYPHEN-MINUS
))


def _validate_tooth_chars(raw: str, field: str) -> None:
    """Raise ``ValueError`` if *raw* contains characters other than 1-8 or A-E."""
    invalid = set(raw) - _VALID_TOOTH_CHARS
    if invalid:
        raise ValueError(
            f"Invalid character(s) in {field}: {', '.join(sorted(invalid))}. "
            f"Only 1-8 and A-E are allowed."
        )


def _build_alt_text(cmd: dict, mode: str) -> str:
    """Build an alt-text description of a Palmer command.

    *mode* must be one of ``"FDI"``, ``"Universal"``, ``"Anatomical"``, or
    ``"Alphanumeric"``.
    Raises ``ValueError`` if any quadrant contains invalid tooth characters.

    Supports two shorthand conventions in addition to bare tooth characters:

    * **Range notation** in quadrant fields: ``1-4`` is equivalent to ``1234``,
      ``A-C`` is equivalent to ``ABC``.
    * **Midline dash**: when the upper/lower midline field is a dash (``-``,
      en-dash, or em-dash), each adjacent quadrant value is interpreted as
      a *count* of teeth from the midline.  E.g. ``upper_mid="-"`` with
      ``UL="2"`` expands UL to ``"12"`` (teeth 1 and 2).
    """
    if mode not in ("FDI", "Universal", "Anatomical", "Alphanumeric"):
        raise ValueError(f"Unknown alt-text mode: {mode!r}")

    # --- Midline analysis ------------------------------------------------
    upper_mid = _strip_tex_commands(cmd.get("upper_mid", ""))
    lower_mid = _strip_tex_commands(cmd.get("lower_mid", ""))
    is_novert = upper_mid == "novert" or lower_mid == "novert"
    upper_is_dash = upper_mid in _MIDLINE_DASHES
    lower_is_dash = lower_mid in _MIDLINE_DASHES

    # A midline dash between a permanent-tooth side and a deciduous-tooth
    # side is not valid Palmer notation.  Skip alt-text entirely.
    if upper_is_dash:
        ul_raw = _strip_tex_commands(cmd.get("UL", ""))
        ur_raw = _strip_tex_commands(cmd.get("UR", ""))
        if ul_raw and ur_raw and ul_raw[0].isdigit() != ur_raw[0].isdigit():
            return ""
    if lower_is_dash:
        ll_raw = _strip_tex_commands(cmd.get("LL", ""))
        lr_raw = _strip_tex_commands(cmd.get("LR", ""))
        if ll_raw and lr_raw and ll_raw[0].isdigit() != lr_raw[0].isdigit():
            return ""

    # --- Expand quadrant fields ------------------------------------------
    quad_raw: dict[str, str] = {}
    for qk in ("UL", "UR", "LR", "LL"):
        raw = _strip_tex_commands(cmd.get(qk, ""))
        is_upper = qk in ("UL", "UR")
        if (is_upper and upper_is_dash) or (not is_upper and lower_is_dash):
            raw = _expand_count_from_midline(raw)
        else:
            raw = _expand_ranges(raw)
        quad_raw[qk] = raw

    # --- Validate and convert each quadrant ------------------------------
    quad_keys = ("UL", "UR", "LR", "LL")
    parts: list[str] = []

    for qk in quad_keys:
        raw = quad_raw[qk]
        if not raw:
            continue
        _validate_tooth_chars(raw, qk)

        if mode == "FDI":
            if is_novert:
                # FDI numbers encode the quadrant (left/right); with novert
                # the side is unknown, so no meaningful FDI notation exists.
                continue
            perm_q, decid_q = _FDI_QUAD[qk]
            teeth: list[str] = []
            for ch in raw:
                if ch in "12345678":
                    teeth.append(str(perm_q * 10 + int(ch)))
                else:  # A-E
                    idx = ord(ch) - ord("A") + 1
                    teeth.append(str(decid_q * 10 + idx))
            parts.append(", ".join(teeth))

        elif mode == "Universal":
            if is_novert:
                # Universal numbers are side-specific; with novert the
                # side is unknown, so no meaningful notation exists.
                continue
            teeth = []
            for ch in raw:
                if ch in "12345678":
                    teeth.append(str(_UNIVERSAL_PERM[qk][ch]))
                else:
                    teeth.append(_UNIVERSAL_DECID[qk][ch])
            parts.append(", ".join(teeth))

        elif mode == "Anatomical":
            if is_novert:
                label = "Maxillary" if qk in ("UL", "UR") else "Mandibular"
            else:
                label = _ANAT_QUAD[qk]
            teeth = []
            for ch in raw:
                name = _ANATOMICAL_NAMES[ch]
                teeth.append(f"{label} {name}")
            parts.append(", ".join(teeth))

        elif mode == "Alphanumeric":
            if is_novert:
                # No left/right distinction: upper → "U", lower → "L".
                prefix = "U" if qk in ("UL", "UR") else "L"
            else:
                # Engine keys use viewer's perspective; Alphanumeric
                # notation uses patient's perspective (L↔R swapped).
                prefix = _ALPHA_QUAD[qk]
            teeth = []
            for ch in raw:
                teeth.append(f"{prefix}{ch}")
            parts.append(", ".join(teeth))

    # --- Midline suffix (non-dash symbols only) --------------------------
    upper_sym = "" if (upper_is_dash or is_novert) else upper_mid
    lower_sym = "" if (lower_is_dash or is_novert) else lower_mid
    midline = ""
    if upper_sym and lower_sym:
        midline = f" (midline: {upper_sym}/{lower_sym})"
    elif upper_sym:
        midline = f" (upper midline: {upper_sym})"
    elif lower_sym:
        midline = f" (lower midline: {lower_sym})"

    if parts:
        return ", ".join(parts) + midline
    return ""


# ---------------------------------------------------------------------------
# Run-level font extraction
# ---------------------------------------------------------------------------

def _para_text(para) -> str:
    """Concatenate all runs in a paragraph into a single string."""
    return "".join(r.text for r in para.runs)


def _run_at_offset(runs, char_offset: int):
    """Return the run object that covers *char_offset* in concatenated text.

    Returns ``None`` if the offset is out of range.
    """
    pos = 0
    for run in runs:
        end = pos + len(run.text)
        if char_offset < end:
            return run
        pos = end
    return None


def _get_para_style_rfonts(run):
    """Return the ``<w:rFonts>`` element from the run's paragraph style, or ``None``.

    Looks up the paragraph style (``<w:pStyle>``) of the paragraph containing
    *run*, then returns that style's ``<w:rPr>/<w:rFonts>`` element so callers
    can read its font attributes.  Only the immediate paragraph style is
    checked — base-style walking is intentionally omitted to keep the
    implementation simple; the existing docDefaults fallback handles styles
    that inherit their fonts from Normal.

    Returns ``None`` when any part of the lookup fails (no style, no rPr,
    no rFonts, or any unexpected exception).
    """
    try:
        para_elem = run._element.getparent()
        if para_elem is None or para_elem.tag != qn('w:p'):
            return None
        pPr = para_elem.find(qn('w:pPr'))
        if pPr is None:
            return None
        pStyle_elem = pPr.find(qn('w:pStyle'))
        if pStyle_elem is None:
            return None
        style_id = pStyle_elem.get(qn('w:val'))
        if not style_id:
            return None
        styles_part_elem = run.part.styles._element  # type: ignore[union-attr]
        for style in styles_part_elem.findall(qn('w:style')):
            if style.get(qn('w:styleId')) == style_id:
                rPr = style.find(qn('w:rPr'))
                if rPr is not None:
                    return rPr.find(qn('w:rFonts'))
                break
    except (AttributeError, KeyError) as exc:
        logger.debug("Para style rFonts lookup failed: %s", exc)
    return None


def _extract_font(run, on_debug: Callable[[str], None] | None = None) -> tuple[str, float, str]:
    r"""Extract font family, size (pt), and text colour from a python-docx Run.

    Returns ``(font_family, font_size_pt, text_color)`` with sensible defaults
    when the run inherits values from its style hierarchy (i.e. the property is
    None).  *text_color* is a ``#RRGGBB`` hex string, or ``""`` when the run
    uses the default (black) colour.

    Both ``w:rFonts/@w:ascii`` and ``w:rFonts/@w:eastAsia`` are read.  When
    the ASCII font is a generic Word default (e.g. Times New Roman, Century)
    **and** the eastAsia font is installed on the system, the eastAsia font is
    preferred — this matches the behaviour users see in Japanese Word documents.
    """
    font_name = DEFAULT_FONT_FAMILY
    font_size_pt = DEFAULT_FONT_SIZE_PT
    text_color = ""

    if run is not None:
        ascii_name = run.font.name              # w:rFonts/@w:ascii
        ascii_explicit = bool(ascii_name)        # True when set directly on the run

        # PRIVATE API: python-docx (>=0.8.11) does not expose w:eastAsia
        # via its public API.  This access pattern is stable across 0.8.x-1.x
        # but may break in future major versions.  Pin python-docx<2.0.
        east_asia_name: str | None = None
        ascii_theme: str | None = None
        east_asia_theme: str | None = None
        try:
            rPr = run._element.rPr
            if rPr is not None and rPr.rFonts is not None:
                east_asia_name = rPr.rFonts.get(qn('w:eastAsia'))
                ascii_theme = rPr.rFonts.get(qn('w:asciiTheme'))
                east_asia_theme = rPr.rFonts.get(qn('w:eastAsiaTheme'))
        except (AttributeError, KeyError):
            pass

        # When the run has no explicit fonts, check the paragraph style first,
        # then fall back to document defaults (w:docDefaults/w:rPrDefault).
        if not ascii_name and not east_asia_name and not ascii_theme and not east_asia_theme:
            style_rfonts = _get_para_style_rfonts(run)
            if style_rfonts is not None:
                ascii_name = ascii_name or style_rfonts.get(qn('w:ascii'))
                east_asia_name = east_asia_name or style_rfonts.get(qn('w:eastAsia'))
                ascii_theme = ascii_theme or style_rfonts.get(qn('w:asciiTheme'))
                east_asia_theme = east_asia_theme or style_rfonts.get(qn('w:eastAsiaTheme'))

        if not ascii_name and not east_asia_name and not ascii_theme and not east_asia_theme:
            try:
                styles_elem = run.part.element.find(
                    './/' + qn('w:docDefaults') + '/' + qn('w:rPrDefault')
                    + '/' + qn('w:rPr') + '/' + qn('w:rFonts'))
                if styles_elem is None:
                    # docDefaults may live in the styles part instead
                    styles_part_elem = run.part.styles._element  # type: ignore[union-attr]
                    styles_elem = styles_part_elem.find(
                        './/' + qn('w:docDefaults') + '/' + qn('w:rPrDefault')
                        + '/' + qn('w:rPr') + '/' + qn('w:rFonts'))
                if styles_elem is not None:
                    ascii_theme = ascii_theme or styles_elem.get(qn('w:asciiTheme'))
                    east_asia_theme = east_asia_theme or styles_elem.get(qn('w:eastAsiaTheme'))
                    # Also pick up literal font names from docDefaults
                    if not ascii_name:
                        ascii_name = styles_elem.get(qn('w:ascii'))
                    if not east_asia_name:
                        east_asia_name = styles_elem.get(qn('w:eastAsia'))
            except (AttributeError, KeyError) as exc:
                logger.debug("docDefaults font lookup failed: %s", exc)

        # Resolve theme references to concrete font names.
        if (ascii_theme or east_asia_theme) and (not ascii_name or not east_asia_name):
            theme = _get_theme_fonts(run)
            if theme:
                if not ascii_name and ascii_theme:
                    # e.g. "minorHAnsi" -> theme["minorHAnsi"]
                    ascii_name = theme.get(ascii_theme)
                if not east_asia_name and east_asia_theme:
                    east_asia_name = theme.get(east_asia_theme)

        _font_msg = (
            f"Font candidates: ascii={ascii_name!r} (theme={ascii_theme!r}), "
            f"eastAsia={east_asia_name!r} (theme={east_asia_theme!r})"
        )
        logger.debug("%s", _font_msg)
        if on_debug:
            on_debug(_font_msg)

        # Prefer the eastAsia font when the ASCII font is absent or an
        # *inherited* generic Word default — this is the typical layout of
        # Japanese (and other CJK) Word documents whose fonts come from
        # docDefaults / theme rather than explicit run-level formatting.
        # When the ASCII font was explicitly set on the run (ascii_explicit),
        # honour the user's choice even if it is a "generic" name like Arial.
        if (
            east_asia_name
            and (not ascii_name
                 or (ascii_name in _GENERIC_ASCII_FONTS and not ascii_explicit))
        ):
            font_name = east_asia_name
        elif ascii_name:
            font_name = ascii_name
        elif east_asia_name:
            font_name = east_asia_name

        _sel_msg = f"Font selected: {font_name!r}"
        logger.debug("%s", _sel_msg)
        if on_debug:
            on_debug(_sel_msg)

        if run.font.size is not None:
            # run.font.size is in EMU; convert to points.
            font_size_pt = run.font.size.pt
        try:
            if run.font.color and run.font.color.rgb:
                text_color = f"#{run.font.color.rgb}"
        except (AttributeError, TypeError):
            pass
    return font_name, font_size_pt, text_color


# ---------------------------------------------------------------------------
# In-place command replacement helpers
# ---------------------------------------------------------------------------

def _find_affected_runs(runs, cmd_start: int, cmd_end: int) -> list[tuple]:
    """Return runs overlapping ``[cmd_start, cmd_end)`` in concatenated text.

    Each entry is ``(run, run_start, run_end)`` where *run_start* / *run_end*
    are character offsets within the concatenated paragraph text.
    """
    affected: list[tuple] = []
    pos = 0
    for run in runs:
        run_start = pos
        run_end = pos + len(run.text)
        if run_end > cmd_start and run_start < cmd_end:
            affected.append((run, run_start, run_end))
        pos = run_end
        if pos >= cmd_end:
            break
    return affected


def _make_text_run_element(src_run_element, text: str):
    """Create a ``<w:r>`` element with the same ``<w:rPr>`` as *src_run_element*
    but with new *text* content.

    This clones the **full** run-properties XML (font, size, bold, italic,
    colour, highlight, strikethrough, superscript, …) so that every formatting
    attribute is preserved — unlike selective attribute copy.
    """
    new_r = OxmlElement("w:r")
    rPr = src_run_element.find(qn("w:rPr"))
    if rPr is not None:
        new_r.append(deepcopy(rPr))
    t_el = OxmlElement("w:t")
    t_el.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_el.set(qn("xml:space"), "preserve")
    new_r.append(t_el)
    return new_r


def _replace_command_inplace(
    para, cmd_start: int, cmd_end: int,
    img_path: Path, dpi: int, alt_text: str,
    valign: str = "base",
    font_size_pt: float = 10.5,
) -> None:
    """Replace a Palmer command at ``[cmd_start, cmd_end)`` with an inline
    image **in-place**, leaving all other runs untouched.

    Only the runs that overlap the command range are modified:

    * Text before the command (in the first affected run) is kept.
    * Text after the command (in the last affected run) is kept.
    * All affected runs in between are removed.
    * A new image run is inserted at the command position.

    When *valign* is ``"center"``, a ``<w:position>`` element is added to
    the image run's ``<w:rPr>`` so that the image is vertically centred on
    the surrounding text.
    """
    runs = list(para.runs)
    affected = _find_affected_runs(runs, cmd_start, cmd_end)
    if not affected:
        return

    p_el = para._element
    first_run, first_start, _first_end = affected[0]
    last_run, last_start, _last_end = affected[-1]

    before_text = first_run.text[: cmd_start - first_start]
    after_text = last_run.text[cmd_end - last_start :]

    # Anchor: the last affected run element — new elements are inserted
    # immediately after it, then the affected runs are cleaned up.
    anchor_el = last_run._element

    # 1) After-text run (clone formatting from the last affected run).
    if after_text:
        after_el = _make_text_run_element(last_run._element, after_text)
        anchor_el.addnext(after_el)

    # 2) Image run — created via python-docx API (appended at end), then
    #    moved to the correct position right after the anchor.
    img_run = para.add_run()
    # Copy full rPr from the first affected run so fallback rendering
    # matches the surrounding text style.
    src_rPr = first_run._element.find(qn("w:rPr"))
    if src_rPr is not None:
        img_run._element.insert(0, deepcopy(src_rPr))

    with Image.open(img_path) as pil_img:
        w_px, h_px = pil_img.size
    w_emu = int(w_px / dpi * EMU_PER_INCH)
    h_emu = int(h_px / dpi * EMU_PER_INCH)
    inline_shape = img_run.add_picture(
        str(img_path), width=Emu(w_emu), height=Emu(h_emu),
    )
    if alt_text:
        inline_shape._inline.docPr.set("descr", alt_text)

    # Vertical alignment: shift the image run down so it centres on the text.
    if valign == "center":
        img_height_pt = h_px / dpi * 72
        # Offset in half-points (negative = lower the content).
        # We want to lower by (img_height - font_size) / 2 points
        # = (img_height - font_size) half-points.
        offset_halfpts = -round(img_height_pt - font_size_pt)
        if offset_halfpts != 0:
            rPr = img_run._element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                img_run._element.insert(0, rPr)
            # Remove any existing <w:position> (may come from copied rPr).
            for old_pos in rPr.findall(qn("w:position")):
                rPr.remove(old_pos)
            pos_el = OxmlElement("w:position")
            pos_el.set(qn("w:val"), str(offset_halfpts))
            rPr.append(pos_el)

    # Move the image run element to right after the anchor.
    img_el = img_run._element
    p_el.remove(img_el)
    anchor_el.addnext(img_el)

    # 3) Clean up affected runs.
    if before_text:
        first_run.text = before_text
    else:
        p_el.remove(first_run._element)

    for run, _, _ in affected[1:]:
        p_el.remove(run._element)


# ---------------------------------------------------------------------------
# Paragraph-level processing
# ---------------------------------------------------------------------------

def _fmt_cmd_detail(
    para_label: str,
    font_name: str,
    font_size_pt: float,
    text_color: str,
    error: Exception | None = None,
) -> str:
    """Build a human-readable detail string for per-command progress logging."""
    parts = [para_label, f"font={font_name}", f"size={font_size_pt}pt"]
    if text_color:
        parts.append(f"color={text_color}")
    detail = ", ".join(parts)
    if error is not None:
        detail += f": {error}"
    return detail


def _process_paragraph(
    para,
    compiler: PalmerCompiler,
    dpi: int,
    tmpdir: Path,
    counter: itertools.count,
    errors: list[str],
    on_command: Callable[[bool, str], None] | None = None,
    para_label: str = "",
    alt_text_mode: str | None = None,
    valign_mode: str = "Force center",
    on_debug: Callable[[str], None] | None = None,
) -> int:
    r"""Find ``\Palmer`` commands in *para*, render them, and replace with
    inline images.  Returns the number of commands successfully replaced.

    *on_command*, when provided, is called once per ``\Palmer`` command with
    ``(success, detail)`` where *detail* is a short description including the
    paragraph label, font name, and error (on failure).

    *alt_text_mode*, when set to ``"FDI"``, ``"Universal"``,
    ``"Anatomical"``, or ``"Palmer command"``, adds a descriptive
    alt-text (``descr`` attribute on ``wp:docPr``) to each inserted image.
    ``None`` (the default) skips alt-text generation.
    """
    full = _para_text(para)
    cmds = find_palmer_commands(full)
    if not cmds:
        return 0

    # Phase 1 — render all Palmer commands and collect results.
    # Commands that fail to render are skipped (original text stays intact).
    # (cmd, img_path, alt_text, valign, font_size_pt)
    rendered: list[tuple[dict, Path, str, str, float]] = []
    replaced = 0

    for cmd in cmds:
        # Determine which run the \Palmer command lives in → extract font & colour.
        src_run = _run_at_offset(para.runs, cmd["start"])
        font_name, font_size_pt, text_color = _extract_font(src_run, on_debug=on_debug)

        try:
            cmd_alt_text = ""
            if alt_text_mode == "Palmer command":
                cmd_alt_text = full[cmd["start"] : cmd["end"]]
            elif alt_text_mode is not None:
                cmd_alt_text = _build_alt_text(cmd, alt_text_mode)

            img = compiler.render(
                UL=cmd.get("UL", ""),
                UR=cmd.get("UR", ""),
                LR=cmd.get("LR", ""),
                LL=cmd.get("LL", ""),
                upper_mid=cmd.get("upper_mid", ""),
                lower_mid=cmd.get("lower_mid", ""),
                option=cmd.get("option") or "base",
                font_family=font_name,
                font_size_pt=font_size_pt,
                text_color=text_color,
                dpi=dpi,
                margin_top=2,
                margin_bottom=2,
                margin_left=2,
                margin_right=2,
                alpha=True,
            )
            out = tmpdir / f"palmer_{next(counter)}.png"
            img.save(str(out), dpi=(dpi, dpi))
            # Determine vertical alignment for this command.
            if valign_mode == "Force center":
                valign = "center"
            else:  # "Follow command option"
                opt = (cmd.get("option") or "base").lower()
                valign = "center" if opt == "center" else "base"
            rendered.append((cmd, out, cmd_alt_text, valign, font_size_pt))
            replaced += 1
            if on_command:
                on_command(True, _fmt_cmd_detail(
                    para_label, font_name, font_size_pt, text_color,
                ))
        except (ValueError, RuntimeError, OSError) as exc:
            # Keep the original text in the document and record the error.
            errors.append(f"Palmer command render failed (font: {font_name}): {exc}")
            if on_command:
                on_command(False, _fmt_cmd_detail(
                    para_label, font_name, font_size_pt, text_color,
                    error=exc,
                ))

    # Phase 2 — replace rendered commands with images **in reverse order**.
    # Processing from back to front ensures that character offsets for
    # earlier commands remain valid after each replacement.
    for cmd, img_path, alt_text, valign, cmd_font_pt in reversed(rendered):
        _replace_command_inplace(
            para, cmd["start"], cmd["end"], img_path, dpi, alt_text,
            valign=valign, font_size_pt=cmd_font_pt,
        )

    return replaced


# ---------------------------------------------------------------------------
# Paragraph collection helpers
# ---------------------------------------------------------------------------

class _PartProxy:
    """Minimal proxy providing ``.part`` for ``Paragraph`` parent resolution.

    When constructing ``Paragraph`` objects from raw ``<w:p>`` elements found
    inside text boxes, the parent must expose a ``.part`` property so that
    ``para.add_run().add_picture()`` can register image relationships.  This
    lightweight adapter wraps an existing ``StoryPart`` and satisfies that
    contract.
    """

    def __init__(self, part):
        self._part = part

    @property
    def part(self):
        return self._part


def _collect_table_paras(
    tables, label_prefix: str = "",
) -> list[tuple]:
    """Collect all paragraphs from *tables* with location labels.

    Handles merged cells by tracking seen underlying XML ``<w:tc>`` elements
    so that the same cell is not processed multiple times — even when
    python-docx returns distinct ``_Cell`` wrappers for the same ``<w:tc>``
    (as happens with vertical merges).  Recurses into nested tables (tables
    inside a cell) so that Palmer commands at any nesting depth are found.
    """
    result: list[tuple] = []
    for t_idx, table in enumerate(tables, 1):
        seen_cells: set[int] = set()
        for r_idx, row in enumerate(table.rows, 1):
            for c_idx, cell in enumerate(row.cells, 1):
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                for p_idx, para in enumerate(cell.paragraphs, 1):
                    label = (
                        f"{label_prefix}table {t_idx} row {r_idx} "
                        f"col {c_idx} para {p_idx}"
                    )
                    result.append((para, label))
                # Recurse into nested tables within this cell.
                if cell.tables:
                    nested_prefix = (
                        f"{label_prefix}table {t_idx} row {r_idx} "
                        f"col {c_idx} > "
                    )
                    result.extend(
                        _collect_table_paras(cell.tables, nested_prefix)
                    )
    return result


def _collect_hf_paras(doc) -> list[tuple]:
    """Collect paragraphs from all headers and footers across sections.

    Skips headers/footers that are linked to the previous section (they share
    the same underlying XML, so processing them again would be a duplicate).
    Guards first-page and even-page variants with the corresponding section
    flags to avoid triggering lazy creation of empty XML parts.
    """
    result: list[tuple] = []
    seen_ids: set[int] = set()

    for s_idx, section in enumerate(doc.sections, 1):
        hf_items: list[tuple] = [
            (section.header, f"section {s_idx} header "),
            (section.footer, f"section {s_idx} footer "),
        ]

        if section.different_first_page_header_footer:
            hf_items.append(
                (section.first_page_header, f"section {s_idx} first-page header ")
            )
            hf_items.append(
                (section.first_page_footer, f"section {s_idx} first-page footer ")
            )

        if hasattr(section, "even_page_header"):
            hf_items.append(
                (section.even_page_header, f"section {s_idx} even-page header ")
            )
        if hasattr(section, "even_page_footer"):
            hf_items.append(
                (section.even_page_footer, f"section {s_idx} even-page footer ")
            )

        for hf, label_prefix in hf_items:
            if hf.is_linked_to_previous:
                continue
            hf_el_id = id(hf._element)
            if hf_el_id in seen_ids:
                continue
            seen_ids.add(hf_el_id)

            for p_idx, para in enumerate(hf.paragraphs, 1):
                result.append((para, f"{label_prefix}para {p_idx}"))

            result.extend(_collect_table_paras(hf.tables, label_prefix))

    return result


def _collect_textbox_paras(doc) -> list[tuple]:
    r"""Collect paragraphs from text boxes by traversing the document XML.

    Text boxes (both modern DrawingML ``<wps:txbx>`` and legacy VML
    ``<v:textbox>``) share the same ``<w:txbxContent>`` wrapper.  This
    function searches every document part (body, headers, footers) for
    those elements and wraps their ``<w:p>`` children as python-docx
    ``Paragraph`` objects.
    """
    result: list[tuple] = []

    # Build a list of (xml_root, story_part) pairs to search.
    parts_to_search: list[tuple] = []

    # Body
    body_part = doc.part
    parts_to_search.append((doc.element, body_part))

    # Headers / footers (only unlinked ones to avoid duplicates)
    seen_hf_ids: set[int] = set()
    for section in doc.sections:
        hf_accessors = ["header", "footer"]
        if section.different_first_page_header_footer:
            hf_accessors += ["first_page_header", "first_page_footer"]

        for attr in hf_accessors:
            hf = getattr(section, attr)
            if hf.is_linked_to_previous:
                continue
            hf_el_id = id(hf._element)
            if hf_el_id in seen_hf_ids:
                continue
            seen_hf_ids.add(hf_el_id)
            parts_to_search.append((hf._element, hf.part))

    seen_txbx: set[int] = set()
    tb_idx = 0
    tag = qn("w:txbxContent")

    for xml_root, story_part in parts_to_search:
        proxy = _PartProxy(story_part)
        for txbx_content in xml_root.iter(tag):
            txbx_id = id(txbx_content)
            if txbx_id in seen_txbx:
                continue
            seen_txbx.add(txbx_id)
            tb_idx += 1

            # Direct paragraphs inside the text box
            for p_idx, p_el in enumerate(txbx_content.findall(qn("w:p")), 1):
                para = Paragraph(p_el, proxy)
                label = f"textbox {tb_idx} para {p_idx}"
                result.append((para, label))

            # Tables inside the text box
            tbl_els = txbx_content.findall(qn("w:tbl"))
            if tbl_els:
                tables = [DocxTable(t_el, proxy) for t_el in tbl_els]
                result.extend(
                    _collect_table_paras(tables, f"textbox {tb_idx} ")
                )

    return result


# ---------------------------------------------------------------------------
# Top-level conversion
# ---------------------------------------------------------------------------

def convert_docx(
    input_path: Path,
    output_path: Path,
    compiler: PalmerCompiler,
    dpi: int = 600,
    on_progress: Callable[[str], None] | None = None,
    alt_text_mode: str | None = None,
    valign_mode: str = "Force center",
    on_debug: Callable[[str], None] | None = None,
    stop_event: threading.Event | None = None,
) -> tuple[int, list[str]]:
    r"""Convert a .docx file by replacing ``\Palmer`` commands with images.

    Args:
        input_path:  Path to the input .docx file.
        output_path: Path to write the output .docx file.
        compiler:    A ``PalmerCompiler`` instance used for rendering.
        dpi:         Resolution for rendered images (default 600).
        on_progress: Optional callback ``(message: str) -> None`` for status
                     updates.  Called from the **calling** thread (not the
                     main GUI thread).
        alt_text_mode: When set to ``"FDI"``, ``"Universal"``,
                     ``"Anatomical"``, or ``"Palmer command"``, each
                     inserted image receives a descriptive alt-text.
                     ``None`` (default) skips alt-text.
        valign_mode: ``"Force center"`` (default) centres every inline
                     image on the text line.  ``"Follow command option"``
                     reads each command's ``[base|center|bottom]`` option
                     and applies centre or baseline alignment accordingly.
        stop_event:  Optional :class:`threading.Event`.  When set, the
                     conversion loop checks it before each paragraph and
                     raises :exc:`ConversionCancelled` if it is set.

    Returns:
        ``(replaced_count, errors)`` where *errors* is a list of error
        message strings for commands that failed to render.

    Raises:
        ConversionCancelled: If *stop_event* is set before all paragraphs
            have been processed.
    """
    def _log(msg: str) -> None:
        if on_progress:
            on_progress(msg)

    input_path = Path(input_path)
    output_path = Path(output_path)

    logger.info(
        "convert_docx: input=%s output=%s dpi=%d backend=%s exe=%s",
        input_path, output_path, dpi, compiler.backend.name, compiler.backend.executable,
    )
    _log(f"Opening {input_path.name} ...")
    doc = Document(str(input_path))

    counter = itertools.count()
    errors: list[str] = []
    replaced = 0

    # Collect paragraphs with location labels from all document areas:
    # body, tables, headers/footers, and text boxes.
    body_paras = list(doc.paragraphs)
    table_paras = _collect_table_paras(doc.tables)
    hf_paras = _collect_hf_paras(doc)
    textbox_paras = _collect_textbox_paras(doc)

    all_labeled_paras = table_paras + hf_paras + textbox_paras

    # Pre-scan: count valid Palmer commands for progress reporting.
    # Uses full parsing via find_palmer_commands() to ensure only complete
    # commands (with all 6 mandatory brace groups) are counted.
    cmd_total = sum(
        len(find_palmer_commands(_para_text(para)))
        for para in itertools.chain(
            body_paras,
            (para for para, _label in all_labeled_paras),
        )
    )

    if cmd_total == 0:
        _log("No Palmer commands found.")
        return 0, errors

    _log(f"Found {cmd_total} Palmer command(s). Rendering ...")
    cmd_done = 0

    def _on_cmd(success: bool, detail: str) -> None:
        nonlocal cmd_done
        cmd_done += 1
        status = "OK" if success else "FAILED"
        _log(f"[{cmd_done}/{cmd_total}] {status} -- {detail}")

    with tempfile.TemporaryDirectory(prefix="palmer_conv_") as tmp:
        tmpdir = Path(tmp)

        # Body paragraphs
        for para_idx, para in enumerate(body_paras, 1):
            if stop_event is not None and stop_event.is_set():
                raise ConversionCancelled()
            label = f"paragraph {para_idx}"
            n = _process_paragraph(
                para, compiler, dpi, tmpdir, counter, errors,
                on_command=_on_cmd, para_label=label,
                alt_text_mode=alt_text_mode,
                valign_mode=valign_mode,
                on_debug=on_debug,
            )
            replaced += n

        # Table cells, headers/footers, text boxes
        for para, label in all_labeled_paras:
            if stop_event is not None and stop_event.is_set():
                raise ConversionCancelled()
            n = _process_paragraph(
                para, compiler, dpi, tmpdir, counter, errors,
                on_command=_on_cmd, para_label=label,
                alt_text_mode=alt_text_mode,
                valign_mode=valign_mode,
                on_debug=on_debug,
            )
            replaced += n

        _log(f"Saving {output_path.name} ...")
        doc.save(str(output_path))

    return replaced, errors
