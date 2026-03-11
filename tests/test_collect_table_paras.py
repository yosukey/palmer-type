"""Tests for _collect_table_paras — paragraph collection from tables.

Covers:
  - Simple table cells
  - Horizontally merged cells (no duplicate paragraphs)
  - Vertically merged cells (no duplicate paragraphs)
  - Combined horizontal + vertical merges
  - Nested tables (paragraphs collected recursively)
  - Nested tables inside merged cells
  - Palmer command detection in table cells
"""

from __future__ import annotations

import sys
from pathlib import Path

# Allow importing from src/ without installing as a package.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from docx import Document

from palmer_converter import _collect_table_paras, find_palmer_commands, _para_text


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_doc_with_simple_table() -> Document:
    """Create a document with a 2x2 table, each cell having one paragraph."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.paragraphs[0].text = f"r{r_idx}c{c_idx}"
    return doc


def _make_doc_with_hmerge() -> Document:
    """Create a document with a table whose first row is horizontally merged."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    # Merge first row: columns 0-1
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).paragraphs[0].text = "merged"
    table.cell(0, 2).paragraphs[0].text = "not merged"
    table.cell(1, 0).paragraphs[0].text = "a"
    table.cell(1, 1).paragraphs[0].text = "b"
    table.cell(1, 2).paragraphs[0].text = "c"
    return doc


def _make_doc_with_vmerge() -> Document:
    """Create a document with a table whose first column is vertically merged."""
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    # Merge first column: rows 0-1
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).paragraphs[0].text = "vmerged"
    table.cell(0, 1).paragraphs[0].text = "r0c1"
    table.cell(1, 1).paragraphs[0].text = "r1c1"
    table.cell(2, 0).paragraphs[0].text = "r2c0"
    table.cell(2, 1).paragraphs[0].text = "r2c1"
    return doc


def _make_doc_with_combined_merge() -> Document:
    """Create a document with both horizontal and vertical merges."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    # Horizontal merge: row 0, cols 0-1
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).paragraphs[0].text = "hmerged"
    table.cell(0, 2).paragraphs[0].text = "r0c2"
    # Vertical merge: rows 1-2, col 2
    table.cell(1, 2).merge(table.cell(2, 2))
    table.cell(1, 2).paragraphs[0].text = "vmerged"
    table.cell(1, 0).paragraphs[0].text = "r1c0"
    table.cell(1, 1).paragraphs[0].text = "r1c1"
    table.cell(2, 0).paragraphs[0].text = "r2c0"
    table.cell(2, 1).paragraphs[0].text = "r2c1"
    return doc


def _make_doc_with_nested_in_merged() -> Document:
    """Create a document with a nested table inside a horizontally merged cell."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    # Horizontal merge: row 0, cols 0-1
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).paragraphs[0].text = "merged_outer"
    # Add a nested table inside the merged cell.
    inner = table.cell(0, 0).add_table(rows=1, cols=2)
    inner.cell(0, 0).paragraphs[0].text = "inner_a"
    inner.cell(0, 1).paragraphs[0].text = "inner_b"
    table.cell(0, 2).paragraphs[0].text = "r0c2"
    table.cell(1, 0).paragraphs[0].text = "r1c0"
    table.cell(1, 1).paragraphs[0].text = "r1c1"
    table.cell(1, 2).paragraphs[0].text = "r1c2"
    return doc


def _make_doc_with_palmer_in_table() -> Document:
    """Create a document with a Palmer command inside a table cell."""
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).paragraphs[0].text = (
        r"Tooth: \Palmer{1}{2}{3}{4}{5}{6}"
    )
    table.cell(1, 0).paragraphs[0].text = "no command here"
    return doc


def _make_doc_with_nested_table() -> Document:
    """Create a document with a table containing a nested table in one cell."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=2)
    outer.cell(0, 0).paragraphs[0].text = "outer_cell"

    # Add a nested table inside cell (0, 1).
    inner_cell = outer.cell(0, 1)
    inner_cell.paragraphs[0].text = "before_nested"
    inner = inner_cell.add_table(rows=1, cols=2)
    inner.cell(0, 0).paragraphs[0].text = "nested_a"
    inner.cell(0, 1).paragraphs[0].text = "nested_b"
    return doc


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_simple_table():
    doc = _make_doc_with_simple_table()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    assert texts == ["r0c0", "r0c1", "r1c0", "r1c1"]

    # Verify labels contain table/row/col info.
    for _para, label in paras:
        assert "table 1" in label


def test_hmerge_no_duplicate():
    doc = _make_doc_with_hmerge()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    # The merged cell should appear only once.
    assert texts.count("merged") == 1
    assert "not merged" in texts
    assert "a" in texts
    assert "b" in texts
    assert "c" in texts


def test_nested_table():
    doc = _make_doc_with_nested_table()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    assert "outer_cell" in texts
    assert "before_nested" in texts
    assert "nested_a" in texts
    assert "nested_b" in texts

    # Verify nested labels contain ">" separator.
    nested_labels = [label for _p, label in paras if "nested_a" in _p.text]
    assert len(nested_labels) == 1
    assert ">" in nested_labels[0]


def test_nested_table_not_duplicated():
    """Nested table paragraphs should not appear as direct children of the
    outer cell (python-docx cell.paragraphs does not include nested table
    content)."""
    doc = _make_doc_with_nested_table()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    # Each text should appear exactly once.
    for t in ("outer_cell", "before_nested", "nested_a", "nested_b"):
        assert texts.count(t) == 1, f"'{t}' appeared {texts.count(t)} times"


def test_vmerge_no_duplicate():
    """Vertically merged cells must not produce duplicate paragraphs.

    python-docx returns distinct _Cell wrappers for the same underlying
    <w:tc> element in continuation rows.  The deduplication must use the
    XML element identity (cell._tc) rather than the Python object identity
    (id(cell)) to avoid processing the same cell twice.
    """
    doc = _make_doc_with_vmerge()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    assert texts.count("vmerged") == 1, (
        f"'vmerged' appeared {texts.count('vmerged')} times (expected 1)"
    )
    assert "r0c1" in texts
    assert "r1c1" in texts
    assert "r2c0" in texts
    assert "r2c1" in texts


def test_combined_hmerge_and_vmerge():
    """Both horizontal and vertical merges in the same table."""
    doc = _make_doc_with_combined_merge()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    assert texts.count("hmerged") == 1, (
        f"'hmerged' appeared {texts.count('hmerged')} times"
    )
    assert texts.count("vmerged") == 1, (
        f"'vmerged' appeared {texts.count('vmerged')} times"
    )
    for expected in ("r0c2", "r1c0", "r1c1", "r2c0", "r2c1"):
        assert expected in texts, f"'{expected}' missing from collected texts"


def test_nested_table_in_merged_cell():
    """Nested table inside a horizontally merged cell — no duplication."""
    doc = _make_doc_with_nested_in_merged()
    paras = _collect_table_paras(doc.tables)

    texts = [p.text for p, _label in paras]
    # Merged outer cell and its nested table should each appear once.
    for expected in ("merged_outer", "inner_a", "inner_b",
                     "r0c2", "r1c0", "r1c1", "r1c2"):
        assert texts.count(expected) == 1, (
            f"'{expected}' appeared {texts.count(expected)} times"
        )

    # Nested table labels should contain the ">" separator.
    nested_labels = [label for _p, label in paras if "inner_a" in _p.text]
    assert len(nested_labels) == 1
    assert ">" in nested_labels[0]


def test_palmer_command_in_table_cell_is_found():
    r"""A \Palmer command inside a table cell is reachable by the
    collection + parsing pipeline."""
    doc = _make_doc_with_palmer_in_table()
    paras = _collect_table_paras(doc.tables)

    found = []
    for para, label in paras:
        cmds = find_palmer_commands(_para_text(para))
        if cmds:
            found.append((label, cmds))

    assert len(found) == 1, f"Expected 1 paragraph with commands, got {len(found)}"
    label, cmds = found[0]
    assert len(cmds) == 1
    assert "table 1" in label
