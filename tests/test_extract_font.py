"""Tests for _extract_font — font extraction from python-docx Runs.

Covers:
  - System font detection via _get_system_fonts / _is_system_font
  - ASCII-only font
  - eastAsia-only font (installed on system)
  - eastAsia preferred over generic ASCII font (CJK document scenario)
  - Non-generic ASCII font preserved even with installed eastAsia
  - Generic ASCII preserved when eastAsia is not installed
  - CJK font names (e.g., Yu Mincho/Yu Gothic aliases)
  - run=None returns defaults
"""

from __future__ import annotations

import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from docx.oxml.ns import qn

import palmer_converter
from palmer_converter import _extract_font, _is_system_font, _get_system_fonts
from palmer_engine import DEFAULT_FONT_FAMILY

# Fake system font set used to mock _get_system_fonts in tests.
_FAKE_SYSTEM_FONTS = frozenset({
    "Times New Roman", "Arial", "Calibri", "Georgia", "Century",
    "Yu Mincho", "MS Mincho", "MS Gothic", "MS PGothic", "Meiryo",
    "游明朝", "游ゴシック", "MS 明朝", "MS ゴシック", "メイリオ",
    "SimSun", "MingLiU", "Malgun Gothic",
})


def _make_run(ascii_font: str | None = None,
              east_asia_font: str | None = None) -> MagicMock:
    """Build a mock python-docx Run with the given rFonts attributes."""
    run = MagicMock()
    run.font.name = ascii_font

    rFonts = MagicMock()
    rFonts.get = MagicMock(side_effect=lambda key: (
        east_asia_font if key == qn('w:eastAsia') else None
    ))

    rPr = MagicMock()
    rPr.rFonts = rFonts if (ascii_font is not None or east_asia_font is not None) else None

    run._element.rPr = rPr
    run.font.size = None
    run.font.color.rgb = None

    return run


def _make_run_no_rpr() -> MagicMock:
    """Run with no rPr element at all."""
    run = MagicMock(spec=[])
    # Re-add only the attributes that _extract_font reads directly.
    run.font = MagicMock()
    run.font.name = None
    run.font.size = None
    run.font.color = MagicMock()
    run.font.color.rgb = None
    run._element = MagicMock()
    run._element.rPr = None
    # `run.part` is intentionally absent (spec=[] blocks unknown attrs)
    # so the docDefaults / theme resolution path raises AttributeError.
    return run


# ---------------------------------------------------------------------------
# _is_system_font / _get_system_fonts tests
# ---------------------------------------------------------------------------

@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_is_system_font_installed(mock_fonts):
    assert _is_system_font("Yu Mincho")
    assert _is_system_font("MS Mincho")
    assert _is_system_font("Meiryo")
    assert _is_system_font("游明朝")
    assert _is_system_font("游ゴシック")


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_is_system_font_not_installed(mock_fonts):
    assert not _is_system_font("NonExistentFont")
    assert not _is_system_font("Segoe UI")


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_is_system_font_empty(mock_fonts):
    assert not _is_system_font("")
    assert not _is_system_font(None)  # type: ignore[arg-type]


@patch.object(palmer_converter, '_get_system_fonts', return_value=frozenset())
def test_is_system_font_fc_list_unavailable(mock_fonts):
    """When fc-list is unavailable (empty set), all fonts are unknown."""
    assert not _is_system_font("Yu Mincho")
    assert not _is_system_font("游明朝")


# ---------------------------------------------------------------------------
# _extract_font tests
# ---------------------------------------------------------------------------

@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_ascii_only(mock_fonts):
    """ASCII font set, no eastAsia -> returns ASCII font."""
    run = _make_run(ascii_font="Arial")
    font, _, _ = _extract_font(run)
    assert font == "Arial"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_no_fonts(mock_fonts):
    """Neither ASCII nor eastAsia set -> returns default."""
    run = _make_run_no_rpr()
    font, _, _ = _extract_font(run)
    assert font == DEFAULT_FONT_FAMILY


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_east_asia_only(mock_fonts):
    """No ASCII font, eastAsia installed -> returns eastAsia."""
    run = _make_run(east_asia_font="MS Mincho")
    font, _, _ = _extract_font(run)
    assert font == "MS Mincho"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_generic_ascii_preserved(mock_fonts):
    """Explicit run-level ASCII='Times New Roman' + eastAsia -> keeps ASCII."""
    run = _make_run(ascii_font="Times New Roman", east_asia_font="Yu Mincho")
    font, _, _ = _extract_font(run)
    assert font == "Times New Roman"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_century_preserved(mock_fonts):
    """Explicit run-level ASCII='Century' + eastAsia -> keeps Century."""
    run = _make_run(ascii_font="Century", east_asia_font="MS Gothic")
    font, _, _ = _extract_font(run)
    assert font == "Century"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_non_generic_ascii_preserved(mock_fonts):
    """ASCII='Georgia' (non-generic), eastAsia='MS Mincho' -> keeps Georgia."""
    run = _make_run(ascii_font="Georgia", east_asia_font="MS Mincho")
    font, _, _ = _extract_font(run)
    assert font == "Georgia"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_generic_ascii_with_unknown_east_asia(mock_fonts):
    """Explicit ASCII='Times New Roman' + unknown eastAsia -> keeps ASCII."""
    run = _make_run(ascii_font="Times New Roman", east_asia_font="Segoe UI")
    font, _, _ = _extract_font(run)
    assert font == "Times New Roman"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_calibri_preserved(mock_fonts):
    """Explicit ASCII='Calibri' + eastAsia='Meiryo' -> keeps Calibri."""
    run = _make_run(ascii_font="Calibri", east_asia_font="Meiryo")
    font, _, _ = _extract_font(run)
    assert font == "Calibri"


# --- CJK font name tests (the original bug) ---

@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_century_with_yu_mincho(mock_fonts):
    """Explicit ASCII='Century' + eastAsia='游明朝' -> keeps Century."""
    run = _make_run(ascii_font="Century", east_asia_font="游明朝")
    font, _, _ = _extract_font(run)
    assert font == "Century"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_times_with_yu_gothic(mock_fonts):
    """Explicit ASCII='Times New Roman' + eastAsia='游ゴシック' -> keeps Times."""
    run = _make_run(ascii_font="Times New Roman", east_asia_font="游ゴシック")
    font, _, _ = _extract_font(run)
    assert font == "Times New Roman"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_century_with_ms_mincho(mock_fonts):
    """Explicit ASCII='Century' + eastAsia='MS 明朝' -> keeps Century."""
    run = _make_run(ascii_font="Century", east_asia_font="MS 明朝")
    font, _, _ = _extract_font(run)
    assert font == "Century"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_arial_with_meiryo(mock_fonts):
    """Explicit ASCII='Arial' + eastAsia='メイリオ' -> keeps Arial."""
    run = _make_run(ascii_font="Arial", east_asia_font="メイリオ")
    font, _, _ = _extract_font(run)
    assert font == "Arial"


def test_run_none():
    """run=None -> returns defaults."""
    font, size, color = _extract_font(None)
    assert font == DEFAULT_FONT_FAMILY
    assert color == ""


# --- fc-list unavailable scenario ---

@patch.object(palmer_converter, '_get_system_fonts', return_value=frozenset())
def test_fc_list_unavailable_explicit_ascii_preserved(mock_fonts):
    """When fc-list is unavailable, explicit ASCII is preserved."""
    run = _make_run(ascii_font="Times New Roman", east_asia_font="游明朝")
    font, _, _ = _extract_font(run)
    assert font == "Times New Roman"


# --- Theme font resolution with real docx ---

def test_theme_font_resolved_from_docx():
    """Runs with no explicit font in a 游明朝-themed docx resolve to 游明朝."""
    from palmer_converter import find_palmer_commands, _run_at_offset, _para_text

    docx_path = Path(__file__).resolve().parent / "test.docx"
    if not docx_path.exists():
        import pytest
        pytest.skip("test.docx not available")

    from docx import Document
    doc = Document(str(docx_path))

    # Paragraph 1 (0-indexed): "b \Palmer{A}{B}{D}{123}{}{}"
    # has no explicit rFonts — font comes from theme (游明朝).
    para = doc.paragraphs[1]
    text = _para_text(para)
    cmds = find_palmer_commands(text)
    assert cmds, f"No Palmer commands found in paragraph 1: {text!r}"
    src_run = _run_at_offset(para.runs, cmds[0]["start"])
    font, _, _ = _extract_font(src_run)
    assert font == "游明朝", f"Expected 游明朝, got {font!r}"


# ---------------------------------------------------------------------------
# Paragraph-style font resolution tests (_get_para_style_rfonts)
# ---------------------------------------------------------------------------

def _make_run_with_para_style(
    *,
    style_id: str = "TableText",
    style_ascii: str | None = None,
    style_east_asia: str | None = None,
    style_ascii_theme: str | None = None,
    style_east_asia_theme: str | None = None,
    run_ascii: str | None = None,
    run_east_asia: str | None = None,
) -> MagicMock:
    """Build a mock Run whose font comes from a paragraph style.

    Creates a minimal lxml XML tree that looks like::

        <w:p>
          <w:pPr><w:pStyle w:val="<style_id>"/></w:pPr>
          <w:r/>   ← the run element
        </w:p>

    The paragraph style definition lives in a fake styles part that contains::

        <w:styles>
          <w:style w:styleId="<style_id>">
            <w:rPr><w:rFonts .../></w:rPr>
          </w:style>
        </w:styles>

    If *run_ascii* or *run_east_asia* is provided, that attribute is set
    directly on the run's ``<w:rFonts>`` to simulate an explicit run-level
    override.
    """
    from lxml import etree

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def wtag(name: str) -> str:
        return f"{{{W_NS}}}{name}"

    # --- paragraph + run element ---
    p_el = etree.Element(wtag("p"))
    pPr = etree.SubElement(p_el, wtag("pPr"))
    pStyle_el = etree.SubElement(pPr, wtag("pStyle"))
    pStyle_el.set(wtag("val"), style_id)
    r_el = etree.SubElement(p_el, wtag("r"))

    # Optional explicit run-level rFonts (simulates a run that partly
    # overrides the style).
    if run_ascii is not None or run_east_asia is not None:
        rPr_el = etree.SubElement(r_el, wtag("rPr"))
        rFonts_el = etree.SubElement(rPr_el, wtag("rFonts"))
        if run_ascii is not None:
            rFonts_el.set(wtag("ascii"), run_ascii)
        if run_east_asia is not None:
            rFonts_el.set(wtag("eastAsia"), run_east_asia)

    # --- styles part element ---
    styles_el = etree.Element(wtag("styles"))
    style_el = etree.SubElement(styles_el, wtag("style"))
    style_el.set(wtag("styleId"), style_id)
    if (style_ascii or style_east_asia
            or style_ascii_theme or style_east_asia_theme):
        style_rPr = etree.SubElement(style_el, wtag("rPr"))
        style_rFonts = etree.SubElement(style_rPr, wtag("rFonts"))
        if style_ascii:
            style_rFonts.set(wtag("ascii"), style_ascii)
        if style_east_asia:
            style_rFonts.set(wtag("eastAsia"), style_east_asia)
        if style_ascii_theme:
            style_rFonts.set(wtag("asciiTheme"), style_ascii_theme)
        if style_east_asia_theme:
            style_rFonts.set(wtag("eastAsiaTheme"), style_east_asia_theme)

    # --- assemble mock run ---
    # Keep run._element as a MagicMock so Python attributes (like .rPr) can be
    # freely set.  Wire .getparent() to return the real lxml <w:p> so that
    # _get_para_style_rfonts() can navigate the paragraph structure.
    run = MagicMock()
    run._element = MagicMock()
    run._element.getparent.return_value = p_el

    # run.font.name: None unless explicit run-level ascii is set
    run.font.name = run_ascii
    run.font.size = None
    run.font.color.rgb = None

    # run._element.rPr: read for run-level rFonts
    if run_ascii is not None or run_east_asia is not None:
        rPr_mock = MagicMock()
        rFonts_mock = MagicMock()
        rFonts_mock.get = MagicMock(side_effect=lambda key: (
            run_ascii if key == qn('w:ascii') else
            run_east_asia if key == qn('w:eastAsia') else
            None
        ))
        rPr_mock.rFonts = rFonts_mock
        run._element.rPr = rPr_mock
    else:
        run._element.rPr = None

    # run.part.styles._element: the fake styles XML (real lxml element so
    # that findall() / find() / get() work correctly in _get_para_style_rfonts).
    styles_mock = MagicMock()
    styles_mock._element = styles_el
    part_mock = MagicMock()
    part_mock.styles = styles_mock
    # element.find(...) for docDefaults → return None (nothing there)
    part_mock.element.find = MagicMock(return_value=None)
    # rels: no theme relationship → theme lookup returns {}
    part_mock.rels = {}
    run.part = part_mock

    return run


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_para_style_ascii_only(mock_fonts):
    """Run with no explicit font; paragraph style has ascii='Calibri'."""
    run = _make_run_with_para_style(style_ascii="Calibri")
    font, _, _ = _extract_font(run)
    assert font == "Calibri"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_para_style_east_asia_preferred_over_generic_ascii(mock_fonts):
    """Para style ascii='Calibri' (generic) + eastAsia='游明朝' → 游明朝."""
    run = _make_run_with_para_style(style_ascii="Calibri", style_east_asia="游明朝")
    font, _, _ = _extract_font(run)
    assert font == "游明朝"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_para_style_non_generic_ascii_preserved(mock_fonts):
    """Para style ascii='Georgia' (non-generic) + eastAsia='游明朝' → Georgia."""
    run = _make_run_with_para_style(style_ascii="Georgia", style_east_asia="游明朝")
    font, _, _ = _extract_font(run)
    assert font == "Georgia"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_run_explicit_font_overrides_para_style(mock_fonts):
    """Explicit run-level font takes priority over paragraph style."""
    run = _make_run_with_para_style(
        style_ascii="Calibri",
        style_east_asia="游明朝",
        run_ascii="Arial",
    )
    font, _, _ = _extract_font(run)
    assert font == "Arial"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_para_style_east_asia_only(mock_fonts):
    """Para style has eastAsia only → eastAsia returned."""
    run = _make_run_with_para_style(style_east_asia="MS Mincho")
    font, _, _ = _extract_font(run)
    assert font == "MS Mincho"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_para_style_no_rfonts_falls_to_default(mock_fonts):
    """Para style exists but has no rPr/rFonts → falls through to default."""
    # style_id is set but no font attrs → style has no <w:rPr><w:rFonts>
    run = _make_run_with_para_style()  # no style_ascii / style_east_asia
    font, _, _ = _extract_font(run)
    assert font == DEFAULT_FONT_FAMILY


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_explicit_arial_with_ms_mincho_on_run(mock_fonts):
    """Explicit run-level ASCII='Arial' + eastAsia='ＭＳ 明朝' -> keeps Arial.

    Reproduces the bug: table cells with explicit Arial were rendered as MS Mincho
    because Arial was in _GENERIC_ASCII_FONTS.
    """
    run = _make_run(ascii_font="Arial", east_asia_font="ＭＳ 明朝")
    font, _, _ = _extract_font(run)
    assert font == "Arial"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_inherited_generic_ascii_prefers_east_asia(mock_fonts):
    """Inherited (style-level) generic ASCII + eastAsia → eastAsia preferred."""
    run = _make_run_with_para_style(
        style_ascii="Times New Roman",
        style_east_asia="Yu Mincho",
    )
    font, _, _ = _extract_font(run)
    assert font == "Yu Mincho"


@patch.object(palmer_converter, '_get_system_fonts', return_value=_FAKE_SYSTEM_FONTS)
def test_inherited_arial_prefers_east_asia(mock_fonts):
    """Inherited (style-level) Arial + eastAsia → eastAsia preferred."""
    run = _make_run_with_para_style(
        style_ascii="Arial",
        style_east_asia="ＭＳ 明朝",
    )
    font, _, _ = _extract_font(run)
    assert font == "ＭＳ 明朝"


def test_explicit_font_not_overridden_by_theme():
    """Runs with explicit rFonts are not overridden by theme fonts."""
    from palmer_converter import find_palmer_commands, _run_at_offset, _para_text

    docx_path = Path(__file__).resolve().parent / "test.docx"
    if not docx_path.exists():
        import pytest
        pytest.skip("test.docx not available")

    from docx import Document
    doc = Document(str(docx_path))

    # Paragraph 0: "a \Palmer{A}{}{}{}{}{}" with explicit ＭＳ 明朝
    para = doc.paragraphs[0]
    text = _para_text(para)
    cmds = find_palmer_commands(text)
    assert cmds
    src_run = _run_at_offset(para.runs, cmds[0]["start"])
    font, _, _ = _extract_font(src_run)
    assert font == "ＭＳ 明朝", f"Expected ＭＳ 明朝, got {font!r}"


def test_table_cell_explicit_arial():
    """Table cell with explicit Arial + eastAsia=ＭＳ 明朝 → Arial.

    Regression test for the bug where table cells with explicitly set Arial
    were incorrectly rendered as ＭＳ 明朝.
    """
    from palmer_converter import (
        find_palmer_commands, _run_at_offset, _para_text,
        _collect_table_paras,
    )

    docx_path = Path(__file__).resolve().parent / "test.docx"
    if not docx_path.exists():
        import pytest
        pytest.skip("test.docx not available")

    from docx import Document
    doc = Document(str(docx_path))

    # Table 1 Row 1 Col 2 Para 1: explicit Arial
    table = doc.tables[0]
    cell = table.rows[0].cells[1]
    para = cell.paragraphs[0]
    text = _para_text(para)
    cmds = find_palmer_commands(text)
    assert cmds, f"No Palmer commands in table 1 row 1 col 2 para 1: {text!r}"
    src_run = _run_at_offset(para.runs, cmds[0]["start"])
    font, _, _ = _extract_font(src_run)
    assert font == "Arial", f"Expected Arial, got {font!r}"

    # Nested: Table 1 Row 1 Col 3 > Table 1 Row 1 Col 2 Para 1: explicit Arial
    nested_table = table.rows[0].cells[2].tables[0]
    ncell = nested_table.rows[0].cells[1]
    npara = ncell.paragraphs[0]
    ntext = _para_text(npara)
    ncmds = find_palmer_commands(ntext)
    assert ncmds, f"No Palmer commands in nested table cell: {ntext!r}"
    nsrc_run = _run_at_offset(npara.runs, ncmds[0]["start"])
    nfont, _, _ = _extract_font(nsrc_run)
    assert nfont == "Arial", f"Expected Arial, got {nfont!r}"
